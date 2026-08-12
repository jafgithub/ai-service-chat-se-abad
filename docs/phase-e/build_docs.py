"""
Turn the Markdown sources in this folder into .docx files for Google Drive.

The documents are written as Markdown so they stay readable and reviewable in git. This script renders them to Word, which Google Drive converts to a Google
Doc cleanly and, unlike Markdown, keeps the diagrams embedded.

Setup (once):
    uv venv .venv && uv pip install --python .venv/bin/python python-docx pillow

Then, from this folder:
    python3 render_diagrams.py     # diagrams first, the docs embed them
    .venv/bin/python build_docs.py

Supported Markdown: headings, paragraphs, bullet and numbered lists, tables,
blockquote callouts, fenced code, images, and inline **bold** / *italic* /
`code`. That is deliberately everything these documents use and nothing more.
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

HERE = Path(__file__).resolve().parent
BUILD = HERE / "build"

INK = RGBColor(0x14, 0x21, 0x3A)
MUTED = RGBColor(0x5B, 0x6B, 0x86)
ACCENT = RGBColor(0x22, 0x58, 0xD4)
CODE_BG = "F2F5F9"
CALLOUT_BG = "F6F8FB"

BODY_FONT = "Calibri"
MONO_FONT = "Consolas"
PAGE_WIDTH_IN = 6.5


def shade(element, fill: str) -> None:
    tag = OxmlElement("w:shd")
    tag.set(qn("w:val"), "clear")
    tag.set(qn("w:fill"), fill)
    element.append(tag)


def left_bar(paragraph, colour: str) -> None:
    borders = OxmlElement("w:pBdr")
    bar = OxmlElement("w:left")
    bar.set(qn("w:val"), "single")
    bar.set(qn("w:sz"), "18")
    bar.set(qn("w:space"), "8")
    bar.set(qn("w:color"), colour)
    borders.append(bar)
    paragraph._p.get_or_add_pPr().append(borders)


def setup_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, colour, before in (
        ("Heading 1", 20, INK, 20), ("Heading 2", 15, ACCENT, 16),
        ("Heading 3", 12.5, INK, 12),
    ):
        st = doc.styles[name]
        st.font.name = BODY_FONT
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = colour
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(6)
        st.paragraph_format.keep_with_next = True


def add_runs(paragraph, text: str, base_size=None) -> None:
    """Render inline **bold**, *italic* and `code`."""
    for part in re.split(r"(\*\*.+?\*\*|(?<!\*)\*(?!\*).+?(?<!\*)\*(?!\*)|`[^`]+`)", text):
        if not part:
            continue
        run = paragraph.add_run()
        if part.startswith("**") and part.endswith("**"):
            run.text, run.bold = part[2:-2], True
        elif part.startswith("`") and part.endswith("`"):
            run.text = part[1:-1]
            run.font.name = MONO_FONT
            run.font.size = Pt((base_size or 10.5) - 1)
            run.font.color.rgb = RGBColor(0x9E, 0x25, 0x5C)
        elif part.startswith("*") and part.endswith("*"):
            run.text, run.italic = part[1:-1], True
        else:
            run.text = part
        if base_size and not run.font.size:
            run.font.size = Pt(base_size)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    header, body = rows[0], rows[1:]
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for cell, text in zip(table.rows[0].cells, header):
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        add_runs(p, f"**{text}**", base_size=9.5)
        shade(cell._tc.get_or_add_tcPr(), "E8F0FF")

    for line in body:
        cells = table.add_row().cells
        for cell, text in zip(cells, line):
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            add_runs(p, text, base_size=9.5)
    doc.add_paragraph()


def add_code(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Inches(0.12)
    shade(p._p.get_or_add_pPr(), CODE_BG)
    run = p.add_run("\n".join(lines))
    run.font.name = MONO_FONT
    run.font.size = Pt(8.8)
    run.font.color.rgb = RGBColor(0x24, 0x2C, 0x3D)


def add_callout(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Inches(0.14)
    shade(p._p.get_or_add_pPr(), CALLOUT_BG)
    left_bar(p, "2258D4")
    add_runs(p, " ".join(lines), base_size=10)


def add_image(doc: Document, path: Path, caption: str) -> None:
    if not path.exists():
        print(f"  ! missing image {path.name}, skipped")
        return
    doc.add_picture(str(path), width=Inches(PAGE_WIDTH_IN))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        run = p.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED


def render(md: str, doc: Document) -> None:
    lines = md.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("```"):
            block, i = [], i + 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            add_code(doc, block)
            i += 1
            continue

        if stripped == "---":
            doc.add_page_break()
            i += 1
            continue

        image = re.match(r"!\[(.*?)\]\((.+?)\)", stripped)
        if image:
            add_image(doc, HERE / image.group(2), image.group(1))
            i += 1
            continue

        if stripped.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{2,}:?", c) for c in cells):
                    rows.append(cells)
                i += 1
            if rows:
                add_table(doc, rows)
            continue

        if stripped.startswith(">"):
            block = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                block.append(lines[i].strip().lstrip(">").strip())
                i += 1
            add_callout(doc, block)
            continue

        heading = re.match(r"(#{1,3})\s+(.*)", stripped)
        if heading:
            level = len(heading.group(1))
            p = doc.add_paragraph(style=f"Heading {level}")
            add_runs(p, heading.group(2))
            i += 1
            continue

        bullet = re.match(r"[-*]\s+(.*)", stripped)
        if bullet:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            add_runs(p, bullet.group(1))
            i += 1
            continue

        number = re.match(r"\d+\.\s+(.*)", stripped)
        if number:
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(3)
            add_runs(p, number.group(1))
            i += 1
            continue

        # Plain paragraph: join the wrapped source lines back together.
        block = []
        while i < len(lines) and lines[i].strip() and not re.match(
                r"(#{1,3}\s|[-*]\s|\d+\.\s|\||>|```|!\[|---$)", lines[i].strip()):
            block.append(lines[i].strip())
            i += 1
        add_runs(doc.add_paragraph(), " ".join(block))


def cover(doc: Document, title: str, subtitle: str, meta: list[str]) -> None:
    doc.add_paragraph().paragraph_format.space_after = Pt(40)

    p = doc.add_paragraph()
    run = p.add_run("AI-Order  |  Conversational Shopping Assistant")
    run.font.size = Pt(11)
    run.bold = True
    run.font.color.rgb = ACCENT
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.size = Pt(27)
    run.bold = True
    run.font.color.rgb = INK
    p.paragraph_format.space_after = Pt(8)

    p = doc.add_paragraph()
    run = p.add_run(subtitle)
    run.font.size = Pt(12.5)
    run.font.color.rgb = MUTED
    p.paragraph_format.space_after = Pt(26)

    for entry in meta:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        add_runs(p, entry, base_size=10)

    doc.add_page_break()


DOCS = [
    ("01_Summary", "Search Speed: Summary",
     "Milestone 1, in-memory catalog index. Measured results."),
    ("02_How_It_Works", "Search Speed: How It Works",
     "The search, one message end to end, and start-up."),
    ("03_What_We_Did", "Search Speed: What We Did",
     "The code change, and every command that was run."),
    ("04_Results", "Search Speed: Results",
     "Measurements, live database comparison, and acceptance."),
    ("05_Technical_Reference", "Search Speed: Technical Reference",
     "For whoever maintains the code next."),
]

# The client asked for no em or en dashes anywhere. Catching it here means a
# stray one can't reach a document unnoticed.
# Written as escapes on purpose: a search-and-replace sweep over this file would
# otherwise rewrite the very characters the check is looking for.
BANNED = {"\u2014": "em dash", "\u2013": "en dash"}


def check_dashes(name: str, text: str) -> list[str]:
    problems = []
    for line_no, line in enumerate(text.splitlines(), 1):
        for ch, label in BANNED.items():
            if ch in line:
                problems.append(f"{name}:{line_no} contains an {label}: {line.strip()[:70]}")
    return problems


def main() -> None:
    BUILD.mkdir(parents=True, exist_ok=True)
    made = []

    problems = []
    for slug, _, _ in DOCS:
        src = HERE / f"{slug}.md"
        if src.exists():
            problems += check_dashes(src.name, src.read_text())
    if problems:
        print("\n".join(problems))
        sys.exit(f"\n{len(problems)} banned dash(es) found. Fix them and re-run.")

    for slug, title, subtitle in DOCS:
        src = HERE / f"{slug}.md"
        if not src.exists():
            print(f"! no source for {slug}.md, skipped")
            continue

        doc = Document()
        setup_styles(doc)
        for section in doc.sections:
            section.left_margin = section.right_margin = Inches(1.0)
            section.top_margin = section.bottom_margin = Inches(0.9)

        cover(doc, title, subtitle, [
            "**Prepared by:** Abad Naseer",
            "**Project:** AI-Order (SmartMarket) conversational shopping assistant",
            "**Environment:** Development, dev.agent.fordev.fun",
            "**Catalog size at time of measurement:** 25,631 products",
        ])
        print(f"building {slug}.docx")
        render(src.read_text(), doc)

        out = BUILD / f"{slug}.docx"
        doc.save(out)
        made.append(out)
        print(f"  -> {out.name}  ({out.stat().st_size / 1024:.0f} KB)")

    if not made:
        sys.exit("Nothing built, are the .md sources present?")
    print(f"\n{len(made)} documents in {BUILD}")


if __name__ == "__main__":
    main()
